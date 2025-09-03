#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Transforma a saída estruturada do Copilot (Markdown) + código-fonte
em um documento Word (.docx) estilizado a partir de um template.

Uso:
  python doc_to_word.py --md saida_copilot.md --src caminho/do/script.py \
                        --template template.docx --out docs/MeuScript.docx
"""
from __future__ import annotations
import re
from pathlib import Path
from typing import Optional, Dict

import typer
from rich import print
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt

app = typer.Typer(add_completion=False)

SECTION_KEYS = [
    "Título", "Descrição", "Entradas", "Saídas",
    "Fluxo de Execução", "Dependências", "Erros Comuns", "Exemplo de Uso"
]

def read_text(p: Path) -> str:
    return p.read_text(encoding="utf-8", errors="ignore")

def split_sections(md_text: str) -> Dict[str, str]:
    """
    Espera blocos '### <Seção>:' conforme instruções globais.
    Retorna dict {secao: conteudo}.
    """
    # normaliza \r\n
    t = md_text.replace("\r\n", "\n")
    # pega tudo a partir do primeiro ### 
    chunks = re.split(r"\n?###\s+", t)
    sections: Dict[str, str] = {}
    for ch in chunks:
        if not ch.strip():
            continue
        # esperado: "Título:\nconteúdo..."
        m = re.match(r"([^:\n]+):\s*\n(.*)", ch, flags=re.S)
        if not m:
            # tolera linha única "Título: algo" também
            m = re.match(r"([^:\n]+):\s*(.*)", ch, flags=re.S)
        if m:
            key = m.group(1).strip()
            val = (m.group(2) or "").strip()
            sections[key] = val
    return sections

def add_code_block(doc: Document, text: str, style_name: str = "CodeBlock"):
    """
    Insere bloco de código (texto preservando quebras).
    Usa estilo 'CodeBlock' se existir; senão, usa estilo normal.
    """
    # tenta estilo; se não existir, cai no normal
    try:
        _ = doc.styles[style_name]
        has_style = True
    except KeyError:
        has_style = False

    lines = text.replace("\r\n", "\n").split("\n")
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        if has_style:
            p.style = style_name
        run = p.add_run(line if line else " ")
        if i < len(lines) - 1:
            run.add_break(WD_BREAK.LINE)

def extract_fenced_code(md: str) -> Optional[str]:
    """
    Retorna o primeiro bloco de código com fences ```...```.
    """
    m = re.search(r"```[^\n]*\n(.*?)```", md, flags=re.S)
    return m.group(1).rstrip("\n") if m else None

@app.command()
def build(
    md: Path = typer.Option(..., help="Arquivo .md com a saída do Copilot"),
    src: Path = typer.Option(..., help="Arquivo de código-fonte relacionado"),
    template: Path = typer.Option(..., help="Modelo do Word (.docx) com estilos"),
    out: Path = typer.Option(..., help="Caminho do .docx de saída"),
    include_source: bool = typer.Option(True, help="Incluir código-fonte completo no final"),
):
    if not md.exists():
        raise typer.BadParameter(f"MD não encontrado: {md}")
    if not src.exists():
        raise typer.BadParameter(f"Fonte não encontrada: {src}")
    if not template.exists():
        raise typer.BadParameter(f"Template não encontrado: {template}")

    md_text = read_text(md)
    sections = split_sections(md_text)
    code_from_md = extract_fenced_code(md_text)
    source_code = read_text(src)

    doc = Document(str(template))

    # Título
    title = sections.get("Título") or src.name
    try:
        doc.add_paragraph(title).style = "DocTitle"
    except Exception:
        doc.add_heading(title, level=1)

    # Seções na ordem predefinida
    for key in SECTION_KEYS:
        content = sections.get(key)
        if not content:
            continue
        # Cabeçalho
        doc.add_heading(key, level=2)
        # Conteúdo: tenta detectar lista vs bloco normal
        if key in {"Entradas", "Saídas", "Dependências", "Erros Comuns"}:
            # Quebra por linhas que começam com "- "
            items = [ln[2:].strip() for ln in content.splitlines() if ln.strip().startswith("- ")]
            if items:
                for it in items:
                    p = doc.add_paragraph(style="List Paragraph")
                    p.add_run("• ").bold = True
                    p.add_run(it)
            else:
                doc.add_paragraph(content, style="NormalText")
        elif key == "Exemplo de Uso" and extract_fenced_code(content):
            add_code_block(doc, extract_fenced_code(content) or "", style_name="CodeBlock")
        else:
            doc.add_paragraph(content, style="NormalText")

    # Se houver código no bloco do exemplo e ainda não inserimos:
    if "Exemplo de Uso" in sections and not extract_fenced_code(sections["Exemplo de Uso"]) and code_from_md:
        doc.add_heading("Exemplo de Uso", level=2)
        add_code_block(doc, code_from_md, style_name="CodeBlock")

    # Anexar código-fonte completo (opcional)
    if include_source:
        doc.add_heading("Código-Fonte (Anexo)", level=2)
        add_code_block(doc, source_code, style_name="CodeBlock")

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"[green]OK[/green] Documento gerado em: [bold]{out}[/bold]")

if __name__ == "__main__":
    app()
